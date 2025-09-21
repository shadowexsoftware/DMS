# src/routers/export/routes.py
from __future__ import annotations
from flask import Blueprint, request, send_file, jsonify
from io import BytesIO
from datetime import datetime
import re
import html
import httpx

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

bp = Blueprint("export", __name__, url_prefix="/export")

# ---------- SVG support (optional via cairosvg) ----------
try:
    import cairosvg
    _HAVE_CAIROSVG = True
except Exception:
    _HAVE_CAIROSVG = False

def _looks_like_svg_bytes(b: bytes) -> bool:
    """Простая эвристика: содержимое похоже на SVG."""
    if not b:
        return False
    head = b[:4096].lstrip()
    if head.startswith(b'<?xml'):
        i = head.find(b'>')
        if i != -1:
            head = head[i + 1:].lstrip()
    return head.lower().startswith(b'<svg')

def _is_svg_by_url_or_ct(url: str, ctype: str | None) -> bool:
    if (url or "").lower().split("?")[0].endswith(".svg"):
        return True
    if ctype:
        ct = ctype.lower()
        if "image/svg+xml" in ct or ct.endswith("/svg+xml"):
            return True
    return False

def _svg_to_png_bytes(svg_bytes: bytes) -> bytes:
    """SVG → PNG. Если cairosvg недоступен, вернём исходник (Word не вставит SVG)."""
    if not _HAVE_CAIROSVG:
        return svg_bytes
    try:
        # Ширина PNG — чтобы влезало в страницу; при необходимости подстрой.
        return cairosvg.svg2png(bytestring=svg_bytes, unsafe=True, output_width=1600)
    except Exception:
        return svg_bytes

# -------- utils: sanitize / parse --------
_SCRIPT_TAG_RE = re.compile(r"<\s*script\b[^>]*>(.*?)<\s*/\s*script\s*>", re.I | re.S)
_STYLE_TAG_RE  = re.compile(r"<\s*style\b[^>]*>(.*?)<\s*/\s*style\s*>", re.I | re.S)
_TAG_RE        = re.compile(r"<[^>]+>")
_BR_RE         = re.compile(r"<\s*br\s*/?\s*>", re.I)
_LI_RE         = re.compile(r"<\s*li\b[^>]*>", re.I)
_ENDLI_RE      = re.compile(r"</\s*li\s*>", re.I)
_P_RE          = re.compile(r"</\s*p\s*>", re.I)

_IMG_RE        = re.compile(r"<img[^>]+src=['\"]([^'\" >]+)['\"][^>]*>", re.I)

def _strip_html_to_text(s: str) -> str:
    """
    Очень простая конверсия HTML → читаемый текст:
    <br> → \\n, </p> → \\n\\n, <li> → "• ", </li> → \\n, остальное удаляем.
    """
    if not isinstance(s, str) or not s.strip():
        return ""

    s = _SCRIPT_TAG_RE.sub("", s)
    s = _STYLE_TAG_RE.sub("", s)

    s = _BR_RE.sub("\n", s)
    s = _P_RE.sub("\n\n", s)
    s = _LI_RE.sub("• ", s)
    s = _ENDLI_RE.sub("\n", s)

    # убираем остальные теги
    s = _TAG_RE.sub("", s)
    # html entity → текст
    s = html.unescape(s)
    # прибираем чуть-чуть пустых строк
    lines = [ln.rstrip() for ln in s.splitlines()]
    out = []
    last_empty = False
    for ln in lines:
        if ln.strip():
            out.append(ln)
            last_empty = False
        else:
            if not last_empty:
                out.append("")
            last_empty = True
    return "\n".join(out).strip()

def _extract_img_urls(s: str) -> list[str]:
    if not isinstance(s, str):
        return []
    return _IMG_RE.findall(s) or []

def _fetch_image(url: str, timeout: float = 20.0) -> BytesIO | None:
    """
    Скачивает ресурс. Если это SVG — конвертирует в PNG.
    Возвращает BytesIO, готовый для docx.add_picture().
    """
    try:
        with httpx.Client(timeout=timeout, follow_redirects=True) as cl:
            r = cl.get(url)
            if r.status_code != 200:
                return None
            ctype = (r.headers.get("content-type") or "").lower()
            content = r.content or b""

            # SVG?
            if _is_svg_by_url_or_ct(url, ctype) or _looks_like_svg_bytes(content):
                png = _svg_to_png_bytes(content)
                bio = BytesIO(png)
                bio.seek(0)
                return bio

            # Иначе — принимаем как есть (jpeg/png/gif/webp/…)
            bio = BytesIO(content)
            bio.seek(0)
            return bio
    except Exception:
        return None

# -------- DOCX builder --------
def _apply_base_styles(doc: Document):
    # Базовый шрифт
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    st.font.size = Pt(11)

    # Заголовок H1
    if "Heading1Custom" not in doc.styles:
        h1 = doc.styles.add_style("Heading1Custom", 1)
        h1.base_style = doc.styles["Heading 1"]
        h1.font.size = Pt(16)

    # Заголовок H2
    if "Heading2Custom" not in doc.styles:
        h2 = doc.styles.add_style("Heading2Custom", 1)
        h2.base_style = doc.styles["Heading 2"]
        h2.font.size = Pt(13)

def _add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    p.style = "Heading1Custom"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def _add_kv_table(doc: Document, data: list[tuple[str, str]]):
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light List"
    for k, v in data:
        row = tbl.add_row()
        row.cells[0].text = k
        row.cells[1].text = v or "—"

def _add_section(doc: Document, title: str, body_html: str):
    doc.add_paragraph().add_run()  # небольшой отступ
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    p.style = "Heading2Custom"

    text = _strip_html_to_text(body_html or "")
    if text:
        for part in text.split("\n"):
            doc.add_paragraph(part)
    else:
        doc.add_paragraph("(нет данных)").italic = True

    # Вставим картинки из HTML, если есть
    urls = _extract_img_urls(body_html or "")
    for u in urls[:6]:  # ограничимся, чтобы не раздуть документ
        bio = _fetch_image(u)
        if bio:
            try:
                doc.add_picture(bio, width=Inches(5.5))
            except Exception:
                # неподдерживаемый формат — пропустим
                pass

def _build_docx(payload: dict) -> BytesIO:
    code = (payload.get("code") or "").strip()
    model = payload.get("model") or {}
    header = payload.get("header") or {}
    sections = payload.get("sections") or {}
    image_url = payload.get("imageUrl") or payload.get("image_url")

    model_label = model.get("label") or ""

    name       = header.get("name") or ""
    material   = header.get("materialCode") or ""
    man_hour   = str(header.get("manHour") or "")
    torque     = str(header.get("torque") or "")
    torque_deg = str(header.get("torqueDegree") or "")

    # Создаём документ
    doc = Document()
    _apply_base_styles(doc)

    _add_title(doc, "Документация по детали")

    # «чипы» — представим в одну строку
    chips = [f"Модель: {model_label}", f"Код: {code}"]
    if material:
        chips.append(f"Материал: {material}")
    doc.add_paragraph(" · ".join(chips))

    # Таблица ключевых полей
    _add_kv_table(doc, [
        ("Наименование", name),
        ("Код детали", material),
        ("Рабочие часы", man_hour),
        ("Крутящий момент", torque),
        ("Степень крутящего момента", torque_deg),
    ])

    # Общая картинка узла (explosive)
    if image_url:
        bio = _fetch_image(image_url)
        if bio:
            doc.add_paragraph()
            doc.add_paragraph("Изображение")
            try:
                doc.add_picture(bio, width=Inches(6.2))
            except Exception:
                doc.add_paragraph("(не удалось вставить изображение)").italic = True

    # Секции
    _add_section(doc, "Описание функции",      sections.get("functionDesc"))
    _add_section(doc, "Технические данные",    sections.get("technical"))
    _add_section(doc, "Электрическая схема",   sections.get("circuit"))
    _add_section(doc, "Шаги разборки и сборки", sections.get("steps"))
    _add_section(doc, "Информация о деталях",  sections.get("part"))

    # футер-пометка
    doc.add_paragraph()
    doc.add_paragraph(f"Сформировано автоматически · {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# -------- route --------
@bp.post("/part")
def export_part_docx():
    """
    Вход:
    {
      "code": "STRUCTURE_CODE",
      "model":  { "label": "X01-X01 (2022款)", ... },
      "header": { "name": "...", "materialCode": "...", "manHour": "...", "torque": "...", "torqueDegree": "..." },
      "sections": {
         "functionDesc": "<html...>",
         "technical": "<html...>",
         "steps": "<html...>",
         "part": "<html...>",
         "circuit": "<html...>"
      },
      "imageUrl": "https://.../explosive.svg"   // опционально
    }
    """
    js = request.get_json(silent=True) or {}
    code = (js.get("code") or "").strip()
    if not code:
        return jsonify({"ok": False, "error": "code is required"}), 400

    try:
        buf = _build_docx(js)
    except Exception as e:
        return jsonify({"ok": False, "error": f"docx build error: {e}"}), 500

    base = (js.get("header", {}).get("materialCode") or code or "document").strip().replace("/", "-")
    stamp = datetime.utcnow().strftime("%Y-%m-%d-%H-%M-%S")
    filename = f"{base}_{stamp}.docx"

    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
        max_age=0,
        conditional=False,
        etag=False
    )
